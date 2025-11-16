import os
import sys
import socket
import time
import argparse
import concurrent.futures
import requests
import json
import logging
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# Optional imports
try:
    import nmap
except Exception:
    nmap = None

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
except Exception:
    webdriver = None

try:
    import sublist3r
    SUBLIST3R_AVAILABLE = True
except Exception:
    sublist3r = None
    SUBLIST3R_AVAILABLE = False

# --- Configuration ---
COMMON_PORTS = [21, 22, 23, 25, 53, 80, 110, 111, 135, 139, 143, 443, 445, 465, 587, 993, 995, 1723, 3306, 3389, 5900, 8080, 8443]
DEFAULT_HEADERS = {"User-Agent": "GhostScan/1.0 (+https://github.com/0xDEADVrm)"}
REPORT_DIR = "reports"
SCREENSHOT_DIR = "screenshots"
LOG_LEVEL = logging.INFO

# --- Logging setup ---
logging.basicConfig(
    level=LOG_LEVEL,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# --- Banners ---
LARGE_BANNER = r"""
                       (                         
                       )\ )                      
 (   (   (       )    (()/(   (                  
 )\  )\  )(     (      /(_)) ))\  (   (    (     
((_)((_)(()\    )\  ' (_))  /((_) )\  )\   )\ )  
\ \ / /  ((_) _((_))  | _ \(_))  ((_)((_) _(_/(  
 \ V /  | '_|| '  \() |   // -_)/ _|/ _ \| ' \)) 
  \_/   |_|  |_|_|_|  |_|_\\___|\__|\___/|_||_|  
                                                 

                0xDEADVrm GhostScan - Automated Recon Toolkit
"""

SMALL_BANNER = r"""
[*] Automated Bug Hunter - Vrm GhostScan
"""

# --- Utility / helpers ---
def safe_mkdir(p):
    try:
        os.makedirs(p, exist_ok=True)
    except Exception as e:
        logging.debug(f"mkdir {p} failed: {e}")

def now_ts():
    return datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

# --- Subdomain enumeration helpers ---
def patched_get_csrftoken(resp_text):
    """
    Robust CSRF extraction for sublist3r (best-effort).
    Accepts HTML/text and returns token string or raises ValueError.
    """
    if not resp_text:
        raise ValueError("Empty response for CSRF extraction")
    # BeautifulSoup attempt
    try:
        soup = BeautifulSoup(resp_text, "html.parser")
        inp = soup.find("input", {"name": "csrfmiddlewaretoken"})
        if inp and inp.has_attr("value"):
            return inp["value"].strip()
    except Exception:
        pass
    # regex fallback
    import re
    m = re.search(r"name=['\"]csrfmiddlewaretoken['\"]\s+value=['\"]([^'\"]+)['\"]", resp_text)
    if m:
        return m.group(1)
    raise ValueError("CSRF token not found")

def enumerate_with_sublist3r(domain, threads=10):
    """
    Try to use sublist3r (patched). Returns list of subdomains or raises.
    """
    if not SUBLIST3R_AVAILABLE:
        raise RuntimeError("sublist3r not available")
    try:
        # attempt patch (different versions may vary)
        try:
            sublist3r.get_csrftoken = patched_get_csrftoken
        except Exception:
            pass
        res = sublist3r.main(domain, threads, None, ports=None, silent=True, verbose=False, enable_bruteforce=False, engines=None)
        if isinstance(res, list):
            return sorted(set(res))
        # older versions may store to internal var or stdout; try best-effort
        try:
            res2 = sublist3r.get_result()
            return sorted(set(res2))
        except Exception:
            raise RuntimeError("sublist3r returned unexpected result")
    except Exception as e:
        raise

def crtsh_subdomains(domain, timeout=10):
    """
    Query crt.sh for subdomains (returns list).
    """
    url = f"https://crt.sh/?q=%25.{domain}&output=json"
    try:
        r = requests.get(url, timeout=timeout, headers=DEFAULT_HEADERS)
        if r.status_code != 200:
            logging.debug(f"crt.sh returned {r.status_code}")
            return []
        data = r.json()
        subs = set()
        for entry in data:
            name_value = entry.get("name_value", "")
            for n in name_value.splitlines():
                n = n.strip()
                if n.endswith(domain):
                    subs.add(n.lower())
        return sorted(subs)
    except Exception as e:
        logging.debug(f"crt.sh error: {e}")
        return []

def brute_force_dns(domain, wordlist=None, max_workers=40, timeout=1.2):
    if wordlist is None:
        wordlist = [
    # Common
    "www", "mail", "webmail", "ftp", "smtp", "imap", "pop", "pop3",
    "cpanel", "cpcalendars", "cpcontacts", "webdisk", "webhost",
    "ns1", "ns2", "ns3", "ns4", "dns", "mx", "autodiscover",
    "m", "mobile", "gateway", "remote", "router",

    # Admin / Login Panels
    "admin", "administrator", "admins", "login", "auth", "secure",
    "sso", "oauth", "signin", "accounts", "users", "dashboard",

    # Development & Staging
    "dev", "development", "test", "testing", "qa", "staging",
    "stage", "beta", "alpha", "preprod", "preview", "demo",
    "internal", "sandbox", "lab", "labs",

    # APIs
    "api", "api2", "api-dev", "api-stage", "graphql", "rest",
    "backend", "backend-dev", "services", "service",

    # Cloud / CDNs
    "cdn", "cdn1", "cdn2", "images", "static", "assets",
    "files", "downloads", "storage", "bucket",

    # Email & Communication
    "email", "mailserver", "smtp2", "mx1", "mx2", "newsletter",
    "lists", "support", "helpdesk", "ticket", "tickets",

    # Security & Monitoring
    "security", "monitor", "monitoring", "uptime", "status",
    "logs", "logging", "metrics", "analytics", "telemetry",

    # E-commerce / Payment
    "shop", "store", "cart", "payment", "payments",
    "billing", "checkout", "gateway",

    # Databases
    "db", "database", "sql", "mysql", "postgres", "mongodb",
    "redis", "cassandra", "elastic", "search",

    # Infrastructure
    "vpn", "vpn1", "vpn2", "proxy", "proxy1", "proxy2",
    "fw", "firewall", "loadbalancer", "lb", "edge",

    # File Sharing
    "fileshare", "share", "ftp2", "ftp-dev", "drop",
    "upload", "uploads", "download", "downloads",

    # IoT / Devices
    "device", "devices", "sensor", "sensors", "iot",
    "cam", "camera", "cams",

    # Educational (useful for colleges)
    "coe", "exam", "exams", "student", "students",
    "staff", "faculty", "portal", "login-portal",

    # CPanel / WHM / Hosting
    "whm", "webmin", "plesk", "siteadmin", "host", "hosting",

    # API Production Variants
    "v1", "v2", "v3", "prod", "production", "live",
    "edge-api", "data", "dataserver",

    # CI/CD
    "git", "github", "gitlab", "bitbucket",
    "jenkins", "ci", "cd", "runner",

    # Cloud Providers
    "aws", "cloud", "azure", "gcp",

    # Misc Attack Surface
    "office", "intranet", "extranet", "gateway",
    "backup", "backups", "old", "legacy", "archive",
    "public", "private", "restricted",

    # High-Value Modern Services
    "sms", "push", "notifications", "websocket", "ws",
    "socket", "stream", "events", "event", "eventhub",
    "bot", "botapi", "botserver",

    # Honeypots / Rare But Important
    "honeypot", "honey", "trap", "decoy"
]
    candidates = [f"{w}.{domain}" for w in wordlist]
    found = set()
    def try_host(h):
        try:
            ip = socket.gethostbyname(h)
            return h
        except Exception:
            return None
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(candidates), max_workers)) as ex:
        futures = {ex.submit(try_host, h): h for h in candidates}
        for fut in concurrent.futures.as_completed(futures):
            res = fut.result()
            if res:
                found.add(res)
    return sorted(found)

def enumerate_subdomains(domain, use_sublist3r=True):
    logging.info(f"Enumerating subdomains for {domain}")
    results = []
    # 1) try sublist3r
    if use_sublist3r and SUBLIST3R_AVAILABLE:
        try:
            logging.info("Trying sublist3r (patched) ...")
            s = enumerate_with_sublist3r(domain, threads=10)
            if s:
                logging.info(f"sublist3r found {len(s)}")
                results.extend(s)
        except Exception as e:
            logging.debug(f"sublist3r failed: {e}")

    # 2) crt.sh
    try:
        logging.info("Querying crt.sh ...")
        s2 = crtsh_subdomains(domain)
        if s2:
            logging.info(f"crt.sh found {len(s2)}")
            results.extend(s2)
    except Exception as e:
        logging.debug(f"crt.sh failed: {e}")

    # 3) DNS brute
    try:
        logging.info("Running small DNS brute-force ...")
        s3 = brute_force_dns(domain)
        if s3:
            logging.info(f"brute-force found {len(s3)}")
            results.extend(s3)
    except Exception as e:
        logging.debug(f"DNS brute failed: {e}")

    # include root domain
    results.append(domain)
    normalized = sorted(set([r.strip().lower() for r in results if r and r.strip()]))
    logging.info(f"Total unique subdomains: {len(normalized)}")
    return normalized

# --- Port scanning ---
def scan_with_nmap(target_ip, ports="1-1024"):
    if nmap is None:
        raise RuntimeError("python-nmap not installed")
    nm = nmap.PortScanner()
    nm.scan(target_ip, ports)
    open_ports = []
    if target_ip in nm.all_hosts():
        for proto in nm[target_ip].all_protocols():
            for p in nm[target_ip][proto].keys():
                try:
                    if nm[target_ip][proto][p]["state"] == "open":
                        open_ports.append(p)
                except Exception:
                    pass
    return sorted(open_ports)

def socket_scan(target_ip, ports=COMMON_PORTS, timeout=0.8, max_workers=60):
    open_ports = []
    def try_port(p):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(timeout)
                res = s.connect_ex((target_ip, p))
                if res == 0:
                    return p
        except Exception:
            return None
        return None
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(ports), max_workers)) as ex:
        futures = {ex.submit(try_port, p): p for p in ports}
        for fut in concurrent.futures.as_completed(futures):
            r = fut.result()
            if r:
                open_ports.append(r)
    return sorted(open_ports)

def scan_ports(target, prefer_nmap=True):
    """
    Accepts hostname or ip. Returns list of open ports (ints).
    """
    logging.info(f"Scanning ports for {target}")
    try:
        ip = socket.gethostbyname(target)
    except Exception:
        ip = target
    # try nmap
    if prefer_nmap and nmap:
        try:
            ports = scan_with_nmap(ip, ports="1-1024")
            if ports:
                logging.info(f"nmap open ports: {ports}")
                return ports
        except Exception as e:
            logging.debug(f"nmap failed: {e}")
    # fallback to socket scan
    try:
        ports = socket_scan(ip, ports=COMMON_PORTS)
        logging.info(f"socket scan open ports: {ports}")
        return ports
    except Exception as e:
        logging.debug(f"socket scan failed: {e}")
    return []

# --- Tech detection ---
def detect_tech(url, timeout=6):
    try:
        r = requests.get(url, timeout=timeout, headers=DEFAULT_HEADERS, verify=False)
        server = r.headers.get("Server", "Unknown")
        xp = r.headers.get("X-Powered-By", "Unknown")
        return server, xp
    except Exception:
        # fallback HTTP if HTTPS failed
        if url.startswith("https://"):
            try:
                return detect_tech("http://" + url[len("https://"):], timeout=timeout)
            except Exception:
                pass
        return "Unknown", "Unknown"

# --- Screenshot (selenium) ---
def take_screenshot_selenium(url, out_path, wait=2):
    if webdriver is None:
        logging.debug("Selenium not available")
        return False
    try:
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--window-size=1280,1024")
        # try to instantiate Chrome/Chromium
        driver = webdriver.Chrome(options=options)
    except Exception as e:
        logging.debug(f"webdriver start failed: {e}")
        return False
    try:
        driver.set_page_load_timeout(15)
        driver.get(url)
        time.sleep(wait)
        safe_mkdir(os.path.dirname(out_path) or ".")
        driver.save_screenshot(out_path)
        logging.info(f"Saved screenshot: {out_path}")
        return True
    except Exception as e:
        logging.debug(f"screenshot failed for {url}: {e}")
        return False
    finally:
        try:
            driver.quit()
        except Exception:
            pass

# --- Reporting ---
def generate_reports(domain, subdomains, port_data, tech_data, screenshots):
    safe_mkdir(REPORT_DIR)
    ts = now_ts()
    html_path = os.path.join(REPORT_DIR, f"{domain}_recon_{ts}.html")
    docx_path = os.path.join(REPORT_DIR, f"{domain}_recon_{ts}.docx")

    # HTML
    html = [
        "<!doctype html><html><head><meta charset='utf-8'><title>Recon Report</title></head><body>",
        f"<h1>Recon Report for {domain}</h1><hr>",
        f"<h2>Subdomains ({len(subdomains)})</h2><ul>"
    ]
    for s in subdomains:
        html.append(f"<li>{s}</li>")
    html.append("</ul><hr><h2>Port Scans</h2><ul>")
    for t, ports in port_data.items():
        html.append(f"<li>{t}: {', '.join(map(str, ports)) if ports else 'None'}</li>")
    html.append("</ul><hr><h2>Tech</h2><ul>")
    for t, tech in tech_data.items():
        html.append(f"<li>{t}: Server: {tech[0]}, X-Powered-By: {tech[1]}</li>")
    html.append("</ul><hr><h2>Screenshots</h2><div>")
    for s in screenshots:
        if os.path.exists(s):
            html.append(f"<div style='display:inline-block;margin:6px'><img src='file://{os.path.abspath(s)}' height='140'><br>{os.path.basename(s)}</div>")
    html.append("</div></body></html>")

    with open(html_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(html))
    logging.info(f"HTML report written: {html_path}")

    # DOCX
    doc = Document()
    doc.add_heading(f"Recon Report for {domain}", level=0)
    doc.add_heading("Subdomains", level=1)
    for s in subdomains:
        doc.add_paragraph(s)
    doc.add_heading("Port Scans", level=1)
    for t, ports in port_data.items():
        doc.add_paragraph(f"{t}: {', '.join(map(str, ports)) if ports else 'None'}")
    doc.add_heading("Tech", level=1)
    for t, tech in tech_data.items():
        doc.add_paragraph(f"{t}: Server: {tech[0]}, X-Powered-By: {tech[1]}")
    doc.add_heading("Screenshots", level=1)
    for s in screenshots:
        try:
            if os.path.exists(s):
                doc.add_picture(s, width=Inches(2))
            else:
                doc.add_paragraph(f"Missing: {s}")
        except Exception as e:
            doc.add_paragraph(f"Could not add screenshot {s}: {e}")
    doc.save(docx_path)
    logging.info(f"DOCX report written: {docx_path}")
    return html_path, docx_path

# --- Main workflow ---
def run_recon(target, no_screenshot=False, no_nmap=False, use_sublist3r=True, fast=False):
    print(LARGE_BANNER)
    logging.info(f"Starting recon for: {target}")

    subdomains = enumerate_subdomains(target, use_sublist3r=use_sublist3r)
    port_data = {}
    tech_data = {}
    screenshots = []

    # Limit concurrency depending on fast flag
    max_workers = 10 if fast else 30

    for sd in subdomains:
        logging.info(f"Processing: {sd}")
        # Ports
        try:
            ports = scan_ports(sd if not sd.replace(".", "").isdigit() else sd, prefer_nmap=(not no_nmap))
        except Exception as e:
            logging.debug(f"scan_ports exception for {sd}: {e}")
            ports = []
        port_data[sd] = ports

        # Tech
        try:
            url_https = f"https://{sd}"
            tech = detect_tech(url_https)
            if tech == ("Unknown", "Unknown") and url_https.startswith("https://"):
                tech = detect_tech("http://" + sd)
        except Exception as e:
            logging.debug(f"detect_tech exception for {sd}: {e}")
            tech = ("Unknown", "Unknown")
        tech_data[sd] = tech

        # Screenshot
        if not no_screenshot:
            safe_mkdir(SCREENSHOT_DIR)
            fname = os.path.join(SCREENSHOT_DIR, f"{sd.replace('/', '_')}.png")
            ok = take_screenshot_selenium(f"https://{sd}", fname)
            if not ok:
                ok = take_screenshot_selenium(f"http://{sd}", fname)
            if ok and os.path.exists(fname):
                screenshots.append(fname)
        else:
            logging.debug("Screenshots disabled by flag")

    # Generate reports
    html_report, docx_report = generate_reports(target, subdomains, port_data, tech_data, screenshots)
    logging.info("Recon complete.")
    return {
        "domain": target,
        "subdomains": subdomains,
        "ports": port_data,
        "tech": tech_data,
        "screenshots": screenshots,
        "html_report": html_report,
        "docx_report": docx_report
    }

# --- CLI ---
def parse_args():
    p = argparse.ArgumentParser(prog="Automated Bug Hunter")
    p.add_argument("target", nargs="?", help="Target domain (example.com)")
    p.add_argument("--no-screenshot", action="store_true", help="Disable screenshots")
    p.add_argument("--no-nmap", action="store_true", help="Disable nmap usage (use socket scan only)")
    p.add_argument("--no-sublist3r", action="store_true", help="Do not attempt sublist3r")
    p.add_argument("--fast", action="store_true", help="Use faster mode (fewer threads / checks)")
    p.add_argument("--verbose", "-v", action="store_true", help="Verbose logging")
    return p.parse_args()

def main():
    args = parse_args()
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    if not args.target:
        args.target = input("Enter the target domain (e.g., example.com): ").strip()
    if not args.target:
        logging.error("No target provided. Exiting.")
        sys.exit(1)

    result = run_recon(
        args.target,
        no_screenshot=args.no_screenshot,
        no_nmap=args.no_nmap,
        use_sublist3r=not args.no_sublist3r,
        fast=args.fast
    )

    logging.info(f"Reports: {result.get('html_report')}, {result.get('docx_report')}")

if __name__ == "__main__":
    # suppress insecure request warnings from urllib3 (only for detect_tech)
    try:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    except Exception:
        pass
    main()
